from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DOCX = BASE_DIR / "requerimientos-funcionales-por-rol-podat.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.size = Pt(11)


def add_heading_paragraph(document, text, level):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    if level == 1:
      run.font.size = Pt(18)
    elif level == 2:
      run.font.size = Pt(14)
    else:
      run.font.size = Pt(12)
    return paragraph


def main():
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Requerimientos Funcionales del Sistema PODAT")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Documento orientado a la experiencia de uso, menús, funcionalidades y restricciones por rol"
    )
    run.italic = True
    run.font.size = Pt(11)

    document.add_paragraph("")

    sections = [
        (
            "1. Objetivo del documento",
            [
                "El presente documento describe los requerimientos funcionales del sistema PODAT desde el punto de vista del usuario que ingresa a la aplicación e interactúa con sus pantallas, menús y módulos. La descripción se centra en lo que cada usuario visualiza, las acciones que puede realizar y las restricciones que tiene según su rol dentro del sistema.",
            ],
        ),
        (
            "2. Objetivo general del sistema",
            [
                "PODAT es una plataforma académica web destinada a la gestión de información institucional vinculada con materias, alumnos, notas, asistencias e indicadores estadísticos. El sistema organiza la experiencia de uso según el perfil del usuario autenticado, permitiendo que cada rol acceda únicamente a las funciones que le corresponden.",
            ],
        ),
        (
            "3. Roles de usuario contemplados",
            [
                "El sistema reconoce los siguientes roles principales:",
            ],
        ),
        (
            "4. Acceso al sistema",
            [
                "Cuando el usuario ingresa a la aplicación visualiza una pantalla inicial con dos opciones de acceso: Acceso Docente y Acceso Administrador. Ambas opciones utilizan autenticación mediante cuenta de Google.",
                "Antes de autenticarse, el usuario debe seleccionar el perfil con el que desea ingresar. Luego de la autenticación, si la sesión y los permisos son válidos, el sistema lo redirige al panel principal correspondiente. Si existe un error de autenticación o autorización, el sistema informa la situación y restringe el acceso.",
            ],
        ),
        (
            "5. Comportamiento general después del ingreso",
            [
                "Una vez autenticado, el usuario accede a un panel interno con navegación lateral. El contenido del menú cambia según el rol del usuario.",
                "Elementos comunes de la experiencia de uso:",
            ],
        ),
        (
            "6. Menú y funcionalidades del Administrador",
            [
                "El usuario con rol Administrador visualiza las siguientes opciones principales en el menú lateral: Dashboard, Perfil, Gestión de usuarios, Importar, Alumnos y Materias.",
            ],
        ),
    ]

    for heading, paragraphs in sections:
        document.add_heading(heading, level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

        if heading == "3. Roles de usuario contemplados":
            add_bullet(document, "Administrador.")
            add_bullet(document, "Docente.")
            document.add_paragraph(
                "Además, puede existir un usuario autenticado sin rol consolidado, aunque el funcionamiento normal del sistema se organiza sobre los roles Administrador y Docente."
            )

        if heading == "5. Comportamiento general después del ingreso":
            add_bullet(document, "Identificación visual del usuario.")
            add_bullet(document, "Acceso al perfil.")
            add_bullet(document, "Acceso al panel principal.")
            add_bullet(document, "Selector de tema visual.")
            add_bullet(document, "Opción para volver al inicio.")
            add_bullet(document, "Opción para cerrar sesión.")

    admin_modules = [
        (
            "6.1 Dashboard",
            [
                "El Administrador puede ingresar al panel principal del sistema y visualizar gráficos e indicadores académicos organizados por materia y año.",
                "Puede consultar estadísticas de todas las materias disponibles, cambiar el año global del panel y analizar múltiples tipos de gráficos.",
                "No presenta restricciones funcionales relevantes dentro de este módulo, salvo la disponibilidad real de datos cargados.",
            ],
        ),
        (
            "6.2 Perfil",
            [
                "En la sección Perfil, el Administrador puede visualizar su información personal y editar datos como nombre visible, teléfono, área o departamento, descripción personal y materias declaradas en el perfil.",
                "Esta pantalla no se utiliza para cambiar roles de otros usuarios.",
            ],
        ),
        (
            "6.3 Gestión de usuarios",
            [
                "Esta opción es exclusiva del Administrador. Permite visualizar el listado de usuarios, buscarlos por correo, consultar la última conexión y cambiar su rol entre Docente, Administrador o Pendiente.",
                "El sistema impide que el único administrador existente se quite a sí mismo ese rol.",
            ],
        ),
        (
            "6.4 Importar",
            [
                "El Administrador puede importar estadísticas académicas desde archivos Excel .xlsx.",
                "Puede cargar el archivo, revisar la previsualización, filtrar filas por estado, detectar cambios contra la base actual, guardar estadísticas nuevas o actualizar existentes y crear materias faltantes detectadas durante la importación.",
                "No se guardan cambios sin revisión previa y los indicadores calculados se detectan pero no se persisten como valores base.",
            ],
        ),
        (
            "6.5 Alumnos",
            [
                "Es un módulo integral que reúne gestión de padrón, carga de notas y carga de asistencias. Antes de operar, el Administrador debe seleccionar materia y año.",
            ],
        ),
        (
            "6.6 Materias",
            [
                "Esta opción es exclusiva del Administrador. Permite crear materias manualmente, importarlas desde Excel, actualizar su código, asignarlas a docentes y eliminar asignaciones existentes.",
            ],
        ),
    ]

    for heading, paragraphs in admin_modules:
        document.add_heading(heading, level=2)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

        if heading == "6.5 Alumnos":
            document.add_heading("6.5.1 Gestión de padrón", level=3)
            document.add_paragraph(
                "Permite importar alumnos desde Excel o cargarlos manualmente, revisar una previsualización antes de guardar y asociarlos a una materia y año determinados."
            )
            document.add_paragraph(
                "Restricciones: no se puede confirmar la importación si faltan datos obligatorios o si la estructura del archivo no cumple con lo esperado."
            )
            document.add_heading("6.5.2 Gestión de notas", level=3)
            document.add_paragraph(
                "Permite seleccionar materia, año, nombre de evaluación y tipo de evaluación; luego cargar la lista del curso, ingresar notas, marcar ausentes y guardar."
            )
            document.add_paragraph(
                "Restricciones: las notas deben estar entre 1 y 10 y los recuperatorios solo se habilitan para alumnos que cumplan las reglas del sistema."
            )
            document.add_heading("6.5.3 Gestión de asistencias", level=3)
            document.add_paragraph(
                "Permite seleccionar materia, año y fecha; registrar el tema de la clase, cargar la lista de alumnos, marcar presente, ausente o justificado y guardar."
            )
            document.add_paragraph(
                "El sistema recalcula la condición académica de asistencia del alumno en función del historial disponible."
            )

    document.add_heading("7. Menú y funcionalidades del Docente", level=1)
    document.add_paragraph(
        "El usuario con rol Docente visualiza las siguientes opciones principales en el menú lateral: Dashboard, Perfil, Importar, Alumnos y Mis Materias."
    )

    docente_modules = [
        (
            "7.1 Dashboard",
            [
                "El Docente puede acceder a un panel principal similar al del Administrador, pero limitado a las materias a las que tiene acceso según sus asignaciones.",
                "Puede consultar estadísticas de sus materias, cambiar la materia seleccionada y analizar datos históricos.",
                "No puede consultar información de materias a las que no fue asignado.",
            ],
        ),
        (
            "7.2 Perfil",
            [
                "El Docente puede ver y editar su información personal, consultar su rol, declarar materias en su perfil y visualizar las materias asignadas oficialmente por la administración.",
                "También dispone de accesos rápidos a Mis Materias, Notas y Asistencias.",
            ],
        ),
        (
            "7.3 Importar",
            [
                "El Docente puede importar estadísticas académicas de sus materias mediante archivos Excel .xlsx.",
                "Puede cargar el archivo, revisar la previsualización, filtrar filas por estado, analizar cambios y confirmar el guardado de estadísticas válidas.",
                "No puede crear materias faltantes desde esta pantalla.",
            ],
        ),
        (
            "7.4 Alumnos",
            [
                "El Docente accede a un espacio de trabajo integrado para gestionar padrón, notas y asistencias de sus cursos asignados. Antes de operar debe seleccionar materia y año.",
            ],
        ),
        (
            "7.5 Mis Materias",
            [
                "Esta pantalla permite consultar las materias asignadas oficialmente, sus códigos y las comisiones asociadas cuando existan.",
                "El Docente no puede crear materias, modificar asignaciones ni reasignarse materias desde esta sección.",
            ],
        ),
    ]

    for heading, paragraphs in docente_modules:
        document.add_heading(heading, level=2)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

        if heading == "7.4 Alumnos":
            document.add_heading("7.4.1 Gestión de padrón", level=3)
            document.add_paragraph(
                "Permite importar alumnos desde Excel o cargarlos manualmente, revisar una previsualización y confirmar la importación sobre la materia y año seleccionados."
            )
            document.add_paragraph(
                "No puede trabajar sobre materias que no le pertenecen ni guardar información sin pasar por la revisión previa."
            )
            document.add_heading("7.4.2 Gestión de notas", level=3)
            document.add_paragraph(
                "Permite seleccionar evaluación y tipo, cargar la lista de alumnos del curso, ingresar notas, marcar ausentes y guardar la evaluación."
            )
            document.add_paragraph(
                "No puede registrar notas fuera del rango permitido ni cargar recuperatorios para alumnos no habilitados por las reglas del sistema."
            )
            document.add_heading("7.4.3 Gestión de asistencias", level=3)
            document.add_paragraph(
                "Permite configurar una clase con fecha y tema, cargar la lista del curso, marcar presente, ausente o justificado y guardar la asistencia."
            )
            document.add_paragraph(
                "No puede operar sobre cursos fuera de sus asignaciones ni omitir los datos mínimos requeridos para registrar una clase."
            )

    document.add_heading("8. Funcionalidades accesibles pero no visibles como menú principal", level=1)
    document.add_paragraph(
        "El sistema también contiene pantallas operativas relacionadas con Notas y Asistencias. Estas pueden alcanzarse desde accesos rápidos o por navegación interna, aunque conceptualmente forman parte del trabajo sobre alumnos y cursos."
    )

    document.add_heading("9. Restricciones generales por rol", level=1)
    document.add_heading("9.1 Restricciones del Docente", level=2)
    docente_restricciones = [
        "No puede ingresar a Gestión de usuarios.",
        "No puede ingresar a Materias.",
        "No puede cambiar roles de usuarios.",
        "No puede asignar materias a otros usuarios.",
        "No puede administrar el catálogo global de materias.",
        "No puede consultar información fuera de las materias a las que tiene acceso.",
    ]
    for item in docente_restricciones:
        add_bullet(document, item)

    document.add_heading("9.2 Restricciones del Administrador", level=2)
    admin_restricciones = [
        "No puede guardar datos inválidos.",
        "No puede omitir las reglas de validación de importación.",
        "No puede quitarse el rol de Administrador si es el único administrador registrado.",
        "Depende de la existencia de datos consistentes para que ciertos módulos muestren información útil.",
    ]
    for item in admin_restricciones:
        add_bullet(document, item)

    document.add_heading("10. Requisitos funcionales principales del sistema", level=1)
    requisitos = [
        "Iniciar sesión con Google seleccionando un perfil de acceso.",
        "Mostrar un menú acorde al rol autenticado.",
        "Restringir opciones no autorizadas según el rol.",
        "Visualizar y editar el perfil del usuario autenticado.",
        "Consultar materias según alcance del rol.",
        "Gestionar alumnos por materia y año.",
        "Cargar notas por evaluación.",
        "Cargar asistencias por clase.",
        "Importar estadísticas desde archivos Excel.",
        "Visualizar dashboards y gráficos académicos.",
        "Mostrar mensajes de error, información y confirmación durante cada proceso.",
        "Redirigir al usuario cuando intenta acceder a una función que no le corresponde.",
    ]
    for item in requisitos:
        add_bullet(document, item)

    document.add_heading("11. Matriz resumen de permisos", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Rol", "Puede hacer", "No puede hacer"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")

    rows = [
        (
            "Administrador",
            "Acceder al Dashboard; editar su Perfil; gestionar usuarios y roles; importar estadísticas; gestionar alumnos, padrón, notas y asistencias; crear materias; importar materias; asignar materias a docentes; ver información global del sistema.",
            "No puede guardar datos inválidos ni omitir validaciones. No puede quitarse el rol de administrador si es el único administrador registrado.",
        ),
        (
            "Docente",
            "Acceder al Dashboard de sus materias; editar su Perfil; importar estadísticas de sus materias; gestionar alumnos, padrón, notas y asistencias de sus materias; consultar Mis Materias.",
            "No puede gestionar usuarios; no puede administrar el catálogo global de materias; no puede asignar materias; no puede acceder a funciones exclusivas del Administrador.",
        ),
    ]
    for role, can_do, cannot_do in rows:
        row = table.add_row().cells
        set_cell_text(row[0], role)
        set_cell_text(row[1], can_do)
        set_cell_text(row[2], cannot_do)

    document.add_heading("12. Conclusión", level=1)
    document.add_paragraph(
        "PODAT organiza su funcionamiento en torno a una experiencia diferenciada por rol. El Administrador posee capacidades de gestión global del sistema, mientras que el Docente trabaja sobre sus materias asignadas y los procesos académicos asociados a ellas. Esta separación funcional permite que cada usuario visualice únicamente las opciones pertinentes para su tarea, reduciendo errores operativos y fortaleciendo el control de acceso dentro de la aplicación."
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    section = document.sections[-1]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("PODAT - Requerimientos funcionales por rol")
    footer_run.font.size = Pt(9)

    document.save(OUTPUT_DOCX)
    print(f"Documento generado en: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
